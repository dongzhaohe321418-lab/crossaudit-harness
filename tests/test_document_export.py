"""Final PDF/DOCX delivery is renderable, inspectable, and fail-closed."""
from __future__ import annotations

import io
import zipfile
from pathlib import Path
from dataclasses import replace

import pytest

from crossaudit.auditor.prompt import render_increment
from crossaudit.dcl import run_checks
from crossaudit.document_export import (
    SOURCE_SUFFIX,
    export_instructions,
    extract_document,
    parse_export_task,
    render_docx,
    render_export,
    render_pdf,
    validate_export_work,
)
from crossaudit.errors import ProviderDenial
from crossaudit import generator


SOURCE = """# Verified report

English and 中文混排 are both preserved in the final document.

## Findings

- First auditable point
- 第二个审计要点

| Item | Result |
| --- | --- |
| Accuracy | Passed |
| 中文 | 已通过 |
"""


@pytest.mark.parametrize("format_name,suffix,signature", [
    ("pdf", ".pdf", b"%PDF-"),
    ("docx", ".docx", b"PK"),
])
def test_render_export_yields_only_one_valid_final_binary(
        science: Path, format_name: str, suffix: str, signature: bytes):
    task = "Write the report" + export_instructions(format_name)
    relative = f"experiments/report{SOURCE_SUFFIX}"
    validate_export_work(science, {relative: SOURCE}, task)
    source = science / relative
    source.write_text(SOURCE)

    written = render_export(science, [relative], task)

    assert written == [f"experiments/report{suffix}"]
    assert not source.exists()
    final = science / written[0]
    assert final.read_bytes().startswith(signature)
    view = extract_document(written[0], final.read_bytes())
    assert view.valid and "Verified report" in view.text
    assert "中文" in view.text and len(view.digest) == 64


@pytest.mark.parametrize("format_name", ["pdf", "docx"])
def test_export_transfers_binding_into_same_receipt_scope(
        science: Path, format_name: str):
    task = "Write the report" + export_instructions(format_name)
    relative = f"experiments/scoped{SOURCE_SUFFIX}"
    receipt = generator.apply(
        generator.bind_file_identities(generator.Work("source", {
            relative: SOURCE,
        }), science, ["experiments"]), science, ["experiments"])

    rendered = render_export(science, receipt, task)
    assert rendered is receipt
    final = science / rendered[0]
    assert final.exists() and not (science / relative).exists()

    rendered.rollback()
    assert not final.exists() and not (science / relative).exists()


def test_export_contract_rejects_extra_files_and_unowned_overwrite(science: Path):
    task = "Write" + export_instructions("pdf")
    with pytest.raises(ProviderDenial, match="exactly one"):
        validate_export_work(science, {
            f"experiments/report{SOURCE_SUFFIX}": SOURCE,
            "experiments/notes.txt": "extra",
        }, task)
    target = science / "experiments" / "report.pdf"
    target.write_bytes(b"owner document")
    with pytest.raises(ProviderDenial, match="not previously generated"):
        validate_export_work(science, {
            f"experiments/report{SOURCE_SUFFIX}": SOURCE,
        }, task)


def test_export_marker_is_unambiguous_and_machine_readable():
    task = "Report" + export_instructions("docx")
    request = parse_export_task(task)
    assert request and request.format == "docx" and request.suffix == ".docx"
    with pytest.raises(ProviderDenial, match="conflicting"):
        parse_export_task(task + export_instructions("pdf"))


@pytest.mark.parametrize("path,data", [
    ("experiments/broken.pdf", b"not a PDF"),
    ("experiments/broken.docx", b"not a ZIP"),
])
def test_invalid_office_container_is_a_mandatory_dcl_blocker(path: str, data: bytes):
    result = run_checks({path: data}, [])
    assert result.hard_failures == 1
    finding = result.findings[0].as_dict()
    assert finding["rule"] == "DCL:document-export-integrity"
    assert finding["artifact"] == path


@pytest.mark.parametrize("format_name", ["pdf", "docx"])
def test_auditor_reads_semantics_recovered_from_the_final_binary(
        science: Path, format_name: str):
    task = "Report" + export_instructions(format_name)
    relative = f"experiments/final{SOURCE_SUFFIX}"
    source = science / relative
    source.write_text(SOURCE)
    final_relative = render_export(science, [relative], task)[0]
    final = science / final_relative

    rendered, bounded = render_increment({final_relative: final.read_bytes()})

    assert not bounded
    assert "recovered from final binary" in rendered
    assert "Verified report" in rendered and "中文" in rendered
    assert extract_document(final_relative, final.read_bytes()).digest in rendered


@pytest.mark.parametrize("format_name", ["pdf", "docx"])
def test_build_loop_commits_only_the_final_document(
        science: Path, cfg, format_name: str, monkeypatch):
    from crossaudit import generator as generator_mod
    from crossaudit.cli import build as build_mod
    from crossaudit.controller import StateStore
    from crossaudit.errors import EXIT_ESCALATED

    task = "Write a verified report" + export_instructions(format_name)
    source = f"experiments/loop{SOURCE_SUFFIX}"
    monkeypatch.setattr(build_mod, "_generator_complete", lambda *_a, **_k: object())
    monkeypatch.setattr(build_mod.gen_mod, "generate", lambda **_kwargs:
                        generator_mod.Work("write report", {source: SOURCE}))

    def fake_audit(_args):
        from tests.conftest import git

        sha = git("rev-parse", "HEAD", cwd=science)
        store = StateStore(cfg.root / cfg.state_dir / "state.json")
        cycle = store.open_or_advance(cfg.science_repo, sha, None)
        store.record_verdict(cycle["cycle_id"], sha, "BLOCKED", "receipt", 1)
        return EXIT_ESCALATED

    monkeypatch.setattr(build_mod, "cmd_run", fake_audit)
    monkeypatch.chdir(science)
    scoped = replace(cfg, scope_dirs=["experiments"], max_rounds=1)

    assert build_mod.run_loop(scoped, task) == EXIT_ESCALATED
    final = source.removesuffix(SOURCE_SUFFIX) + "." + format_name
    from tests.conftest import git
    committed = git("show", "--pretty=", "--name-only", "HEAD", cwd=science).splitlines()
    assert final in committed and source not in committed
    assert not (science / source).exists()
    assert extract_document(final, (science / final).read_bytes()).valid


# --------------------------------------------------------------------------
# End-to-end generation correctness, checked against independent libraries so
# a regression in our renderer cannot hide behind our own extractor.
# --------------------------------------------------------------------------
RICH_SOURCE = """# 审计报告 Audit Report

This paragraph mixes English and 中文段落 so the body font must be Unicode.

## Findings 发现

- First auditable point 第一条
- Second point 第二条要点

1. Numbered step one
2. Numbered step 步骤二

| Metric 指标 | Result 结果 |
| --- | --- |
| Accuracy 准确度 | Passed 通过 |
| Coverage 覆盖率 | 95% 达标 |

```python
def audit(value):
    return value >= 0.90  # 阈值检查 threshold
```

Inline `变量` and `ascii_var` are both code spans.

> A block quote 引用 for emphasis.
"""

RICH_PROBES = [
    "Audit Report", "审计报告", "中文段落", "Findings", "发现",
    "第一条", "第二条要点", "步骤二", "准确度", "Passed", "通过",
    "def audit", "阈值检查", "变量", "ascii_var", "引用",
]


def test_pdf_render_is_parseable_by_pypdf_with_full_text_recovered(tmp_path: Path):
    from pypdf import PdfReader

    output = tmp_path / "rich.pdf"
    render_pdf(RICH_SOURCE, output)

    data = output.read_bytes()
    assert data.startswith(b"%PDF-")
    reader = PdfReader(io.BytesIO(data))          # independent of our extractor
    assert not reader.is_encrypted and len(reader.pages) >= 1
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    missing = [probe for probe in RICH_PROBES if probe not in text]
    assert not missing, f"pypdf did not recover: {missing}"


def test_docx_render_is_a_valid_zip_container_with_full_text_recovered(tmp_path: Path):
    from docx import Document

    output = tmp_path / "rich.docx"
    render_docx(RICH_SOURCE, output)

    data = output.read_bytes()
    assert data.startswith(b"PK")
    with zipfile.ZipFile(io.BytesIO(data)) as archive:   # independent container check
        assert "word/document.xml" in archive.namelist()
        assert archive.testzip() is None
    document = Document(io.BytesIO(data))
    values = [p.text for p in document.paragraphs]
    for table in document.tables:
        values += [cell.text for row in table.rows for cell in row.cells]
    text = "\n".join(values)
    missing = [probe for probe in RICH_PROBES if probe not in text]
    assert not missing, f"python-docx did not recover: {missing}"


@pytest.mark.parametrize("render,suffix", [(render_pdf, ".pdf"), (render_docx, ".docx")])
def test_cjk_inside_code_survives_rather_than_becoming_tofu(tmp_path, render, suffix):
    """Regression: Courier carries no CJK, so a non-Latin code run must fall back
    to the Unicode face instead of silently rendering unextractable glyphs."""
    source = ("# Code fidelity\n\n"
              "```python\nx = 1  # 阈值检查\n```\n\nInline `变量` and `ascii_only`.\n")
    output = tmp_path / ("f" + suffix)
    render(source, output)
    view = extract_document(output.as_posix(), output.read_bytes())
    assert view.valid
    for probe in ("阈值检查", "变量", "ascii_only"):
        assert probe in view.text, f"{probe!r} lost in {suffix}"


# --------------------------------------------------------------------------
# Fail-closed: hostile-but-plausible containers must never validate. These are
# structurally well-formed (unlike "not a PDF"), so they probe the real guards.
# --------------------------------------------------------------------------
def _macro_ooxml() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("word/document.xml", "<w:document><w:body/></w:document>")
        archive.writestr("word/vbaProject.bin", b"\x00\x01macro")
    return buffer.getvalue()


def _zip_without_document_part() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("hello.txt", "not a word document")
    return buffer.getvalue()


def _crc_corrupt_docx() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_STORED) as archive:
        archive.writestr("word/document.xml", "<w:document><w:body/></w:document>")
        archive.writestr("word/settings.xml", "X" * 200)
    data = bytearray(buffer.getvalue())
    data[data.find(b"X" * 200) + 40] ^= 0xFF        # flip a stored payload byte
    return bytes(data)


# Compound-file (OLE) magic: a password-protected .docx is encrypted this way and
# is not a ZIP at all, so it must be refused rather than parsed.
_OLE_DOCX = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512


def test_extract_document_is_fail_closed_for_hostile_containers(tmp_path: Path):
    from pypdf import PdfWriter

    plain = tmp_path / "plain.pdf"
    render_pdf("# Confidential\n\nreal body\n", plain)
    writer = PdfWriter(clone_from=str(plain))
    writer.encrypt("secret")
    encrypted = io.BytesIO()
    writer.write(encrypted)

    cases = {
        "encrypted.pdf": (encrypted.getvalue(), "encrypted"),
        "macro.docx": (_macro_ooxml(), "macro"),
        "nopart.docx": (_zip_without_document_part(), "Word document part"),
        "crc.docx": (_crc_corrupt_docx(), "CRC"),
        "ole.docx": (_OLE_DOCX, "not a zip"),
    }
    for name, (data, reason_fragment) in cases.items():
        view = extract_document(name, data)
        assert not view.valid, f"{name} should be rejected"
        assert reason_fragment.lower() in view.reason.lower(), (name, view.reason)


@pytest.mark.parametrize("name,data,fragment", [
    ("experiments/macro.docx", None, "macro"),
    ("experiments/encrypted.pdf", None, "encrypted"),
])
def test_hostile_but_wellformed_documents_are_dcl_blockers(name, data, fragment, tmp_path):
    """A structurally valid container that hides a macro or encryption is a hard
    gate failure, not merely garbage-bytes rejection."""
    from pypdf import PdfWriter

    if name.endswith(".docx"):
        payload = _macro_ooxml()
    else:
        plain = tmp_path / "p.pdf"
        render_pdf("# Body\n\ntext\n", plain)
        writer = PdfWriter(clone_from=str(plain))
        writer.encrypt("pw")
        buffer = io.BytesIO()
        writer.write(buffer)
        payload = buffer.getvalue()

    result = run_checks({name: payload}, [])
    assert result.hard_failures == 1
    finding = result.findings[0].as_dict()
    assert finding["rule"] == "DCL:document-export-integrity"
    assert finding["artifact"] == name
