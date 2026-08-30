"""Auditable, local PDF and DOCX export from a model-authored Markdown source.

The model never emits an opaque binary.  A guided task asks it for exactly one
temporary ``*.crossaudit.md`` source, this module renders that source locally,
validates the resulting container, extracts the final document's own text, and
then removes the source before Git staging.  The committed increment therefore
contains only the file the user requested while the Auditor sees text recovered
from the exact binary named in the receipt manifest.
"""
from __future__ import annotations

import hashlib
import html
import io
import os
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath

from .errors import ProviderDenial
from .file_identity import AppliedFiles, apply_file_payloads
from .gitio import OVERSIZE_DOCUMENT_MARKER

EXPORT_MARKER = "CROSSAUDIT-DOCUMENT-EXPORT"
SOURCE_SUFFIX = ".crossaudit.md"
FORMAT_LABELS = {"pdf": "PDF (.pdf)", "docx": "Word (.docx)"}
INTEGRITY_CONTRACT = (
    "Every PDF/DOCX increment must be a valid, non-encrypted document with at "
    "least one page/paragraph and extractable text. Macro-enabled OOXML is refused."
)

#: Decompression-bomb bounds for the DOCX ZIP container. A few-MB archive whose
#: ``word/document.xml`` declares gigabytes of uncompressed XML would OOM the
#: auditor and the Console preview the instant ``testzip()``/``Document()``
#: decompress it, so the declared sizes in the central directory are checked
#: first — never the compressed payload. Every bound sits far above any real
#: .docx (whose members are a handful of small XML parts plus already-compressed
#: images), so a legitimate document is never rejected.
MAX_DOCX_MEMBERS = 8192
MAX_DOCX_MEMBER_BYTES = 256 * 1024 * 1024          # one member, uncompressed
MAX_DOCX_TOTAL_BYTES = 512 * 1024 * 1024           # all members summed, uncompressed
MAX_DOCX_RATIO = 500                               # summed uncompressed : compressed
#: Below this summed-uncompressed size a high ratio cannot OOM anything, so the
#: ratio test is skipped there to keep small, highly repetitive-but-legitimate
#: documents from tripping it.
_DOCX_RATIO_FLOOR = 8 * 1024 * 1024


@dataclass(frozen=True)
class ExportRequest:
    format: str

    @property
    def suffix(self) -> str:
        return "." + self.format

    @property
    def label(self) -> str:
        return FORMAT_LABELS[self.format]


@dataclass(frozen=True)
class DocumentView:
    format: str
    text: str
    units: int
    digest: str
    valid: bool
    reason: str = ""


@dataclass(frozen=True)
class Block:
    kind: str
    text: str = ""
    level: int = 0
    rows: tuple[tuple[str, ...], ...] = ()


def export_instructions(format_name: str) -> str:
    """A machine-readable task clause plus plain guidance for the Generator."""
    if format_name not in FORMAT_LABELS:
        raise ValueError("unsupported document export format")
    return (
        f"\n\n[{EXPORT_MARKER} format={format_name} source_suffix={SOURCE_SUFFIX}]\n"
        f"- The user-facing deliverable is {FORMAT_LABELS[format_name]}.\n"
        f"- Return exactly one Markdown source file whose path ends in "
        f"{SOURCE_SUFFIX}, for example work/report{SOURCE_SUFFIX}.\n"
        "- Do not return a PDF, DOCX, alternate format, companion file, or notes file.\n"
        "- Use Markdown headings, paragraphs, real lists, block quotes, fenced code, "
        "and pipe tables where they improve the document.\n"
        "- The local controller will render the final binary, validate it, recover "
        "its text for independent audit, and remove the temporary source before commit."
    )


def parse_export_task(task: str) -> ExportRequest | None:
    matches = re.findall(
        rf"\[{EXPORT_MARKER}\s+format=(pdf|docx)\s+source_suffix="
        rf"{re.escape(SOURCE_SUFFIX)}\]", task, flags=re.IGNORECASE)
    if not matches:
        return None
    values = {value.lower() for value in matches}
    if len(values) != 1:
        raise ProviderDenial("the task contains conflicting document export formats")
    return ExportRequest(values.pop())


def _safe_source(relative: str) -> PurePosixPath:
    path = PurePosixPath(relative)
    if (path.is_absolute() or ".." in path.parts or "\\" in relative
            or path.as_posix() != relative or not relative.endswith(SOURCE_SUFFIX)):
        raise ProviderDenial(
            f"document export requires one safe *{SOURCE_SUFFIX} source path")
    return path


def _target_for(source: PurePosixPath, request: ExportRequest) -> PurePosixPath:
    stem = source.as_posix()[:-len(SOURCE_SUFFIX)]
    if not stem or stem.endswith("/"):
        raise ProviderDenial("document export source has no output file name")
    return PurePosixPath(stem + request.suffix)


def _owned_target(root: Path, relative: str) -> bool:
    result = subprocess.run(
        ["git", "log", "-1", "--format=%B", "--", relative], cwd=str(root),
        capture_output=True, text=True)
    return result.returncode == 0 and "CrossAudit-Generator:" in result.stdout


def validate_export_work(root: Path, files: dict[str, str], task: str) -> None:
    """Refuse ambiguous or destructive model output before it touches disk."""
    request = parse_export_task(task)
    if request is None:
        return
    if len(files) != 1:
        raise ProviderDenial(
            f"{request.label} export requires exactly one temporary Markdown source; "
            f"the Generator returned {len(files)} files")
    relative = next(iter(files))
    source = _safe_source(relative)
    source_path = root.joinpath(*source.parts)
    if source_path.exists() or source_path.is_symlink():
        raise ProviderDenial(
            f"temporary export source already exists and will not be overwritten: {relative}")
    target = _target_for(source, request)
    target_path = root.joinpath(*target.parts)
    if (target_path.exists() or target_path.is_symlink()) and not _owned_target(
            root, target.as_posix()):
        raise ProviderDenial(
            f"refusing to overwrite a document not previously generated by CrossAudit: "
            f"{target.as_posix()}")


def _split_cells(line: str) -> tuple[str, ...]:
    value = line.strip().strip("|")
    return tuple(cell.strip() for cell in re.split(r"(?<!\\)\|", value))


def _table_separator(line: str) -> bool:
    cells = _split_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_markdown(source: str) -> list[Block]:
    """A bounded Markdown subset designed for predictable office rendering."""
    lines = source.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[Block] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        fence = re.match(r"^\s*```([^`]*)$", line)
        if fence:
            language = fence.group(1).strip()
            index += 1
            body: list[str] = []
            while index < len(lines) and not re.match(r"^\s*```\s*$", lines[index]):
                body.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ProviderDenial("document Markdown contains an unclosed code fence")
            index += 1
            rendered = language + "\n" + "\n".join(body) if language else "\n".join(body)
            blocks.append(Block("code", rendered))
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*#*\s*$", line)
        if heading:
            blocks.append(Block("heading", heading.group(2).strip(), len(heading.group(1))))
            index += 1
            continue
        if index + 1 < len(lines) and "|" in line and _table_separator(lines[index + 1]):
            rows = [_split_cells(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                row = _split_cells(lines[index])
                if len(row) != len(rows[0]):
                    raise ProviderDenial("document Markdown table has inconsistent columns")
                rows.append(row)
                index += 1
            blocks.append(Block("table", rows=tuple(rows)))
            continue
        bullet = re.match(r"^\s*[-+*]\s+(.+)$", line)
        if bullet:
            blocks.append(Block("bullet", bullet.group(1).strip()))
            index += 1
            continue
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered:
            blocks.append(Block("number", numbered.group(1).strip()))
            index += 1
            continue
        quote = re.match(r"^\s*>\s?(.*)$", line)
        if quote:
            body = [quote.group(1)]
            index += 1
            while index < len(lines):
                continued = re.match(r"^\s*>\s?(.*)$", lines[index])
                if not continued:
                    break
                body.append(continued.group(1))
                index += 1
            blocks.append(Block("quote", " ".join(body).strip()))
            continue
        if re.fullmatch(r"\s*(?:-{3,}|\*{3,}|_{3,})\s*", line):
            blocks.append(Block("rule"))
            index += 1
            continue
        body = [line.strip()]
        index += 1
        while index < len(lines) and lines[index].strip():
            candidate = lines[index]
            if (re.match(r"^(?:#{1,6}\s|\s*```|\s*[-+*]\s+|\s*\d+[.)]\s+|\s*>\s?)", candidate)
                    or re.fullmatch(r"\s*(?:-{3,}|\*{3,}|_{3,})\s*", candidate)
                    or (index + 1 < len(lines) and "|" in candidate
                        and _table_separator(lines[index + 1]))):
                break
            body.append(candidate.strip())
            index += 1
        blocks.append(Block("paragraph", " ".join(body)))
    if not blocks:
        raise ProviderDenial("document export source is empty")
    return blocks


INLINE_TOKEN = re.compile(r"(\[.*?\]\(.*?\)|\*\*.*?\*\*|`.*?`|(?<!\*)\*[^*]+\*)")


def _plain_inline(value: str) -> str:
    value = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    return value.replace("\\|", "|").strip()


def semantic_text(source: str) -> str:
    values: list[str] = []
    for block in parse_markdown(source):
        if block.kind == "table":
            values.extend(" | ".join(_plain_inline(cell) for cell in row)
                          for row in block.rows)
        elif block.kind != "rule":
            values.append(_plain_inline(block.text))
    return "\n".join(value for value in values if value).strip()


def _docx_runs(paragraph, value: str, *, font: str = "Arial Unicode MS") -> None:
    from docx.oxml.ns import qn
    for part in INLINE_TOKEN.split(value):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = (part.startswith("*") and part.endswith("*") and not bold)
        code = part.startswith("`") and part.endswith("`")
        link = re.fullmatch(r"\[([^]]+)\]\(([^)]+)\)", part)
        text = (f"{link.group(1)} ({link.group(2)})" if link else
                part[2:-2] if bold else part[1:-1] if italic or code else part)
        run = paragraph.add_run(text)
        run.bold, run.italic = bold, italic
        run.font.name = "Arial Unicode MS" if code else font
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), run.font.name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), run.font.name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")


def _set_docx_style(style, *, size: float, color: str = "000000", bold=False,
                    before=0, after=0, line=1.0, font="Arial Unicode MS") -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    style.font.name = font
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def _table_geometry(table, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    total = sum(widths)
    table.autofit = False
    props = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd"):
        old = props.find(qn(tag))
        if old is not None:
            props.remove(old)
    width = OxmlElement("w:tblW")
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(total))
    props.append(width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    props.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            tcw = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(value))
            margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(margins)
            for side, amount in (("top", 80), ("bottom", 80),
                                 ("start", 120), ("end", 120)):
                node = margins.find(qn("w:" + side))
                if node is None:
                    node = OxmlElement("w:" + side)
                    margins.append(node)
                node.set(qn("w:w"), str(amount))
                node.set(qn("w:type"), "dxa")


def _column_widths(rows: tuple[tuple[str, ...], ...]) -> list[int]:
    weights = [max(6, min(60, max(len(_plain_inline(row[index])) for row in rows)))
               for index in range(len(rows[0]))]
    total = sum(weights)
    widths = [max(900, round(9360 * weight / total)) for weight in weights]
    difference = 9360 - sum(widths)
    widths[widths.index(max(widths))] += difference
    return widths


def render_docx(source: str, output: Path) -> None:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    blocks = parse_markdown(source)
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    _set_docx_style(doc.styles["Normal"], size=11, after=6, line=1.10)
    _set_docx_style(doc.styles["Title"], size=23, bold=True, after=12)
    for name, size, color, before, after in (
            ("Heading 1", 16, "2E74B5", 16, 8),
            ("Heading 2", 13, "2E74B5", 12, 6),
            ("Heading 3", 12, "1F4D78", 8, 4)):
        _set_docx_style(doc.styles[name], size=size, color=color, bold=True,
                        before=before, after=after)
    for name in ("List Bullet", "List Number"):
        _set_docx_style(doc.styles[name], size=11, after=8, line=1.167)
        doc.styles[name].paragraph_format.left_indent = Inches(0.5)
        doc.styles[name].paragraph_format.first_line_indent = Inches(-0.25)
    if "CrossAudit Code" not in doc.styles:
        style = doc.styles.add_style("CrossAudit Code", WD_STYLE_TYPE.PARAGRAPH)
        _set_docx_style(style, size=9.5, color="1F2937", after=8, line=1.0,
                        font="Arial Unicode MS")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("CrossAudit  |  ")
    label.font.size, label.font.color.rgb = Pt(8), RGBColor.from_string("6B7280")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    first_heading = next((b for b in blocks if b.kind != "rule"), None)
    for block in blocks:
        if block.kind == "heading":
            style = "Title" if block is first_heading and block.level == 1 else (
                "Heading 1" if block.level <= 2 else
                "Heading 2" if block.level == 3 else "Heading 3")
            paragraph = doc.add_paragraph(style=style)
            _docx_runs(paragraph, block.text)
        elif block.kind == "paragraph":
            paragraph = doc.add_paragraph()
            _docx_runs(paragraph, block.text)
        elif block.kind in {"bullet", "number"}:
            paragraph = doc.add_paragraph(style=("List Bullet" if block.kind == "bullet"
                                                  else "List Number"))
            _docx_runs(paragraph, block.text)
        elif block.kind == "quote":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.space_after = Pt(8)
            _docx_runs(paragraph, block.text)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string("4B5563")
            ppr = paragraph._p.get_or_add_pPr()
            border = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            for key, value in (("val", "single"), ("sz", "18"),
                               ("space", "8"), ("color", "9CA3AF")):
                left.set(qn("w:" + key), value)
            border.append(left)
            ppr.append(border)
        elif block.kind == "code":
            paragraph = doc.add_paragraph(style="CrossAudit Code")
            paragraph.add_run(block.text)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F4F7")
            paragraph._p.get_or_add_pPr().append(shading)
        elif block.kind == "rule":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            for key, value in (("val", "single"), ("sz", "6"),
                               ("space", "1"), ("color", "D1D5DB")):
                bottom.set(qn("w:" + key), value)
            border.append(bottom)
            paragraph._p.get_or_add_pPr().append(border)
        elif block.kind == "table":
            table = doc.add_table(rows=len(block.rows), cols=len(block.rows[0]))
            table.style = "Table Grid"
            _table_geometry(table, _column_widths(block.rows))
            for row_index, row in enumerate(block.rows):
                for column_index, value in enumerate(row):
                    cell = table.cell(row_index, column_index)
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(0)
                    _docx_runs(paragraph, value)
                    if row_index == 0:
                        for run in paragraph.runs:
                            run.bold = True
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "F2F4F7")
                        cell._tc.get_or_add_tcPr().append(shading)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)

    doc.core_properties.title = _plain_inline(first_heading.text) if first_heading else output.stem
    doc.core_properties.author = "CrossAudit"
    doc.core_properties.last_modified_by = "CrossAudit"
    doc.save(output)


def _pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    name = "CrossAuditUnicode"
    if name in pdfmetrics.getRegisteredFontNames():
        return name
    candidates = (
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    )
    for candidate in candidates:
        if Path(candidate).is_file():
            try:
                pdfmetrics.registerFont(TTFont(name, candidate))
                pdfmetrics.registerFontFamily(
                    name, normal=name, bold=name, italic=name, boldItalic=name)
                return name
            except Exception:
                pass
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    pdfmetrics.registerFontFamily(
        "STSong-Light", normal="STSong-Light", bold="STSong-Light",
        italic="STSong-Light", boldItalic="STSong-Light")
    return "STSong-Light"


def _pdf_inline(value: str, code_font: str = "Courier") -> str:
    escaped = html.escape(value)
    escaped = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", escaped)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+)\*", r"<i>\1</i>", escaped)
    # Courier carries no CJK glyphs, so monospaced runs holding non-Latin text
    # would render as unextractable tofu and vanish from the audited document.
    # Keep Courier for ASCII code and fall back to the Unicode face otherwise.
    escaped = re.sub(
        r"`([^`]+)`",
        lambda m: (f'<font face="Courier">{m.group(1)}</font>' if m.group(1).isascii()
                   else f'<font face="{code_font}">{m.group(1)}</font>'),
        escaped)
    return escaped.replace("\\|", "|")


def render_pdf(source: str, output: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfgen.canvas import Canvas
    from reportlab.platypus import (HRFlowable, ListFlowable, ListItem, Paragraph,
                                    Preformatted, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)

    blocks = parse_markdown(source)
    font = _pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("CrossAuditBody", parent=styles["BodyText"], fontName=font,
                          fontSize=11, leading=14, spaceAfter=7, textColor=colors.HexColor("#111827"))
    title = ParagraphStyle("CrossAuditTitle", parent=body, fontSize=23, leading=27,
                           spaceAfter=16, textColor=colors.black)
    headings = {
        1: ParagraphStyle("CrossAuditH1", parent=body, fontSize=16, leading=20,
                          spaceBefore=16, spaceAfter=8, textColor=colors.HexColor("#2E74B5")),
        2: ParagraphStyle("CrossAuditH2", parent=body, fontSize=13, leading=17,
                          spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#2E74B5")),
        3: ParagraphStyle("CrossAuditH3", parent=body, fontSize=12, leading=15,
                          spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#1F4D78")),
    }
    quote = ParagraphStyle("CrossAuditQuote", parent=body, leftIndent=18,
                           borderWidth=0, textColor=colors.HexColor("#4B5563"), italic=True)
    code = ParagraphStyle("CrossAuditCode", parent=body, fontName="Courier", fontSize=8.5,
                          leading=11, leftIndent=8, rightIndent=8, borderPadding=8,
                          backColor=colors.HexColor("#F2F4F7"))
    # A code block containing CJK cannot use Courier without dropping those
    # characters; the Unicode face keeps them at the cost of monospacing.
    code_unicode = ParagraphStyle("CrossAuditCodeUnicode", parent=code, fontName=font)
    first_heading = next((b for b in blocks if b.kind != "rule"), None)
    story = []
    list_kind = None
    list_items = []

    def flush_list() -> None:
        nonlocal list_kind, list_items
        if not list_items:
            return
        story.append(ListFlowable(
            [ListItem(Paragraph(_pdf_inline(item, font), body), leftIndent=8)
             for item in list_items],
            bulletType="bullet" if list_kind == "bullet" else "1",
            start="circle" if list_kind == "bullet" else 1,
            leftIndent=24, bulletFontName=font, bulletFontSize=9, spaceAfter=6))
        list_kind, list_items = None, []

    for block in blocks:
        if block.kind in {"bullet", "number"}:
            if list_kind not in (None, block.kind):
                flush_list()
            list_kind = block.kind
            list_items.append(block.text)
            continue
        flush_list()
        if block.kind == "heading":
            chosen = title if block is first_heading and block.level == 1 else headings[min(3, max(1, block.level - 1))]
            story.append(Paragraph(_pdf_inline(block.text, font), chosen))
        elif block.kind == "paragraph":
            story.append(Paragraph(_pdf_inline(block.text, font), body))
        elif block.kind == "quote":
            story.append(Paragraph(_pdf_inline(block.text, font), quote))
        elif block.kind == "code":
            story.append(Preformatted(block.text,
                                      code if block.text.isascii() else code_unicode,
                                      maxLineLength=100))
        elif block.kind == "rule":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#D1D5DB"), spaceBefore=4,
                                    spaceAfter=10))
        elif block.kind == "table":
            rows = [[Paragraph(_pdf_inline(cell, font), body) for cell in row]
                    for row in block.rows]
            widths = [value / 1440 * inch for value in _column_widths(block.rows)]
            table = Table(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D1D5DB")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.extend([table, Spacer(1, 8)])
    flush_list()

    def page(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawRightString(letter[0] - inch, 0.55 * inch,
                               f"CrossAudit  |  {doc.page}")
        canvas.restoreState()

    def canvasmaker(*args, **kwargs):
        kwargs["invariant"] = 1
        return Canvas(*args, **kwargs)

    document = SimpleDocTemplate(str(output), pagesize=letter, rightMargin=inch,
                                 leftMargin=inch, topMargin=inch, bottomMargin=0.82 * inch,
                                 title=_plain_inline(first_heading.text) if first_heading else output.stem,
                                 author="CrossAudit", pageCompression=1)
    document.build(story, onFirstPage=page, onLaterPages=page, canvasmaker=canvasmaker)


def _oversized_archive_reason(archive: zipfile.ZipFile) -> str | None:
    """Reason a DOCX ZIP looks like a decompression bomb, or None if it is sane.

    Reads only the central directory's declared sizes — ``infolist()`` does not
    decompress anything — so the check runs BEFORE ``testzip()``/``Document()``
    detonate the archive. Guards member count, per-member uncompressed size,
    summed uncompressed size, and the overall uncompressed/compressed ratio.
    """
    infos = archive.infolist()
    if len(infos) > MAX_DOCX_MEMBERS:
        return f"{len(infos)} members exceeds the {MAX_DOCX_MEMBERS} limit"
    total_uncompressed = 0
    total_compressed = 0
    for info in infos:
        if info.file_size > MAX_DOCX_MEMBER_BYTES:
            return (f"member {info.filename!r} declares {info.file_size} "
                    f"uncompressed bytes, over the {MAX_DOCX_MEMBER_BYTES} limit")
        total_uncompressed += info.file_size
        total_compressed += info.compress_size
        if total_uncompressed > MAX_DOCX_TOTAL_BYTES:
            return (f"members declare over {MAX_DOCX_TOTAL_BYTES} uncompressed "
                    f"bytes in total")
    ratio = total_uncompressed / max(total_compressed, 1)
    if total_uncompressed > _DOCX_RATIO_FLOOR and ratio > MAX_DOCX_RATIO:
        return (f"uncompressed/compressed ratio {ratio:.0f}:1 exceeds the "
                f"{MAX_DOCX_RATIO}:1 limit")
    return None


def extract_document(path: str, data: bytes) -> DocumentView:
    """Recover semantic text from the exact final bytes used by the Auditor."""
    suffix = Path(path).suffix.lower()
    digest = hashlib.sha256(data).hexdigest()
    if data == OVERSIZE_DOCUMENT_MARKER:
        # gitio refused to buffer a document blob past its guard and handed back
        # this deterministic marker instead. Audit and verify both reach here
        # with the identical bytes, so the document is recorded as too-large on
        # both sides rather than OOMing either of them.
        return DocumentView(suffix.removeprefix(".").upper(), "", 0, digest, False,
                            "document too large to audit")
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            if not data.startswith(b"%PDF-"):
                raise ValueError("missing PDF signature")
            reader = PdfReader(io.BytesIO(data), strict=True)
            if reader.is_encrypted:
                raise ValueError("encrypted PDFs cannot be audited")
            if not reader.pages:
                raise ValueError("PDF has no pages")
            text = "\n\n".join((page.extract_text() or "").strip()
                               for page in reader.pages).strip()
            if not text:
                raise ValueError("PDF has no extractable text")
            return DocumentView("PDF", text, len(reader.pages), digest, True)
        if suffix == ".docx":
            from docx import Document
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                bomb = _oversized_archive_reason(archive)
                if bomb:
                    raise ValueError(
                        f"document too large / suspicious compression: {bomb}")
                names = set(archive.namelist())
                if "word/document.xml" not in names:
                    raise ValueError("DOCX has no Word document part")
                if any(name.lower().endswith("vbaproject.bin") for name in names):
                    raise ValueError("macro-enabled OOXML is not allowed")
                bad = archive.testzip()
                if bad:
                    raise ValueError(f"DOCX ZIP member failed CRC: {bad}")
            document = Document(io.BytesIO(data))
            values = [paragraph.text.strip() for paragraph in document.paragraphs
                      if paragraph.text.strip()]
            for table in document.tables:
                values.extend(" | ".join(cell.text.strip() for cell in row.cells)
                              for row in table.rows)
            text = "\n".join(value for value in values if value).strip()
            if not text:
                raise ValueError("DOCX has no extractable text")
            return DocumentView("DOCX", text, len(document.sections), digest, True)
        return DocumentView(suffix.removeprefix(".").upper(), "", 0, digest, False,
                            "unsupported document format")
    except Exception as exc:
        return DocumentView(suffix.removeprefix(".").upper(), "", 0, digest, False,
                            str(exc)[:300])


def docx_outline(data: bytes) -> list[dict]:
    """Heading/title outline recovered from a DOCX for the Console preview.

    Best-effort and bounded: only paragraphs carrying a Title or ``Heading N``
    style contribute, at most 500 entries, each clamped. Never used by the audit
    path, and only ever called on bytes ``extract_document`` already validated as
    a safe, non-bomb archive, so it re-opens a known-good document.
    """
    from docx import Document
    outline: list[dict] = []
    document = Document(io.BytesIO(data))
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "") or ""
        if style == "Title":
            outline.append({"level": 0, "text": text[:200]})
        else:
            match = re.fullmatch(r"Heading (\d)", style)
            if match:
                outline.append({"level": int(match.group(1)), "text": text[:200]})
        if len(outline) >= 500:
            break
    return outline


def pdf_structure(data: bytes) -> dict:
    """Page count and bookmark outline for the Console PDF preview, best-effort.

    Runs under the same document ceiling the caller enforces and is wrapped by
    the caller in try/except so any parse failure degrades to no metadata rather
    than breaking the sandboxed iframe render.
    """
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data), strict=False)
    if reader.is_encrypted:
        return {}
    pages = len(reader.pages)
    outline: list[dict] = []

    def walk(items, level: int = 1) -> None:
        for item in items:
            if len(outline) >= 500:
                return
            if isinstance(item, list):
                walk(item, level + 1)
                continue
            title = getattr(item, "title", None)
            if title:
                outline.append({"level": level, "text": str(title)[:200]})

    try:
        walk(reader.outline)
    except Exception:
        outline = []
    return {"pages": pages, "outline": outline}


def render_export(root: Path, written: list[str] | AppliedFiles,
                  task: str) -> list[str] | AppliedFiles:
    """Replace the temporary source with one authorized, validated binary."""
    request = parse_export_task(task)
    if request is None:
        return written
    if len(written) != 1:
        raise ProviderDenial("document export received an unexpected file set")
    source = _safe_source(written[0])
    source_path = root.joinpath(*source.parts)
    target = _target_for(source, request)
    source_receipt = written if isinstance(written, AppliedFiles) else None
    target_receipt: AppliedFiles | None = None
    try:
        if source_receipt is not None:
            source_receipt.verify()
            raw = source_receipt.entries[0].payload
        else:
            raw = source_path.read_bytes()
        if not raw:
            raise ProviderDenial("document export source is empty")
        source_text = raw.decode("utf-8")
        expected = semantic_text(source_text)
        if not expected:
            raise ProviderDenial("document export source contains no readable content")
        descriptor, temporary_name = tempfile.mkstemp(
            prefix="crossaudit-export-", suffix=request.suffix)
        os.close(descriptor)
        temporary = Path(temporary_name)
        try:
            if request.format == "pdf":
                render_pdf(source_text, temporary)
            else:
                render_docx(source_text, temporary)
            data = temporary.read_bytes()
            view = extract_document(target.as_posix(), data)
            if not view.valid:
                raise ProviderDenial(
                    f"rendered {request.label} failed validation: {view.reason}")
            expected_terms = [term.casefold() for term in re.findall(r"[\w\u3400-\u9fff]{3,}", expected)]
            recovered = view.text.casefold()
            if expected_terms and sum(term in recovered for term in expected_terms) / len(
                    expected_terms) < 0.90:
                raise ProviderDenial(
                    f"rendered {request.label} did not preserve enough source text")
            target_receipt = apply_file_payloads(
                root, {target.as_posix(): data},
                (list(source_receipt.allowed_dirs)
                 if source_receipt is not None
                 and source_receipt.allowed_dirs is not None else None))
        finally:
            if temporary.exists():
                temporary.unlink()
        if source_receipt is not None:
            # Keep the caller's one lexical receipt scope.  Transferring the
            # derived binary's binding into that object prevents the export
            # seam from returning an unguarded second lifecycle.
            source_receipt.replace_with(target_receipt)
            target_receipt = None
            return source_receipt
        else:
            source_path.unlink()
            # Backward-compatible trusted-call path: callers that supplied a
            # path list have no receipt scope to accept.  Preserve their
            # historical immediate-apply contract explicitly rather than via
            # AppliedFiles.__del__.
            target_receipt.finalize()
            return list(target_receipt)
    except ProviderDenial:
        if target_receipt is not None:
            target_receipt.rollback()
        if source_receipt is not None:
            source_receipt.rollback()
        elif source_path.is_file() and not source_path.is_symlink():
            source_path.unlink()
        raise
    except Exception as exc:
        if target_receipt is not None:
            target_receipt.rollback()
        if source_receipt is not None:
            source_receipt.rollback()
        elif source_path.is_file() and not source_path.is_symlink():
            source_path.unlink()
        raise ProviderDenial(
            f"local {request.label} rendering failed: {type(exc).__name__}: {exc}") from exc


def current_document_text(path: Path) -> str | None:
    if path.suffix.lower() not in {".pdf", ".docx"}:
        return None
    try:
        view = extract_document(path.as_posix(), path.read_bytes())
    except OSError:
        return None
    if not view.valid:
        return None
    return (f"[Rendered {view.format}; the controller recovered this text from the "
            f"current binary for revision.]\n\n{view.text}")
